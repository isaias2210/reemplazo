import os
import re
import sys
from datetime import datetime

from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    flash,
    send_from_directory,
)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =========================
# CONFIGURACIÓN BÁSICA
# =========================

app = Flask(__name__)
app.secret_key = "super-secret-paseu"  # para flash messages

def get_base_dir():
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))

BASE_DIR = get_base_dir()
PLANTILLA_PATH = os.path.join(BASE_DIR, "plantilla.docx")
HISTORIAL_PATH = os.path.join(BASE_DIR, "historial_paseu.csv")
SALIDAS_DIR = os.path.join(BASE_DIR, "salidas")

os.makedirs(SALIDAS_DIR, exist_ok=True)

# =========================
# UTILIDADES
# =========================

def limpiar(s: str) -> str:
    return " ".join(s.replace("\t", " ").split())

def extraer_acudiente(texto: str):
    """
    Busca:
      R.Legal: Nombre Apellido   Cedula: 8-xxx-xxxx
    """
    t = texto.replace("\t", " ")
    m = re.search(
        r"R\.Legal:\s*([A-Za-zÁÉÍÓÚÑ ]+?)\s+Cedula:\s*([0-9-]+)",
        t,
        re.IGNORECASE,
    )
    if m:
        return {
            "NOMBRE_ACUDIENTE": limpiar(m.group(1)).upper(),
            "CEDULA_ACUDIENTE": m.group(2),
        }
    return {"NOMBRE_ACUDIENTE": "", "CEDULA_ACUDIENTE": ""}

def parse_tabla_cheques(texto: str):
    """
    Parsea la tabla tipo:

    Regional    Grado   Centro  ESTUDIANTE  Beca    Monto   N° Cheque ...
    PANAMA NORTE  4 A   ESC. NUEVO PROGRESO  8-1161-843 ALIANIS GUTIERREZ ...

    Devuelve lista de dicts por cheque.
    """
    lineas = [l for l in texto.splitlines() if l.strip()]
    header_idx = None

    # Buscar fila de encabezado
    for i, line in enumerate(lineas):
        if ("Regional" in line and "Grado" in line and "Centro" in line
                and "ESTUDIANTE" in line):
            header_idx = i
            break

    if header_idx is None:
        return []

    filas = []
    for line in lineas[header_idx + 1:]:
        partes = line.split("\t")
        if len(partes) < 13:
            continue

        estudiante = partes[3].strip()  # "4-xxx-xxx NOMBRE APELLIDO"
        tokens = estudiante.split()
        cedula = ""
        nombre = ""

        if tokens and re.match(r"\d{1,2}-\d{3,4}-\d{3,4}", tokens[0]):
            cedula = tokens[0]
            nombre = " ".join(tokens[1:]).upper()
        else:
            nombre = estudiante.upper()

        filas.append({
            "grado": partes[1].strip(),
            "centro": partes[2].strip().upper(),
            "cedula": cedula,
            "nombre": nombre,
            "beca": partes[4].strip(),
            "monto": partes[5].strip(),
            "cheque": partes[6].strip(),
            "planilla": partes[9].strip(),
            "fecha": partes[10].strip(),
            "estado_linea": partes[11].strip(),
            "periodo": partes[12].strip(),
        })

    return filas

def periodo_a_checks(periodo: str):
    periodo = periodo.replace(" ", "")
    d = {"PRIMER_PAGO": "", "SEGUNDO_PAGO": "", "TERCER_PAGO": ""}
    if periodo.startswith("2-"):
        d["PRIMER_PAGO"] = "✓"
    elif periodo.startswith("3-"):
        d["SEGUNDO_PAGO"] = "✓"
    elif periodo.startswith("4-"):
        d["TERCER_PAGO"] = "✓"
    return d

# =========================
# HISTORIAL CSV
# =========================

def guardar_en_historial(datos, telefono, periodo, estado="PENDIENTE"):
    """
    FECHA, NOMBRE_ESTUDIANTE, CEDULA_ESTUDIANTE, TELEFONO, ESTADO,
    NOMBRE_ACUDIENTE, CEDULA_ACUDIENTE, PLANILLA, CHEQUE, NIVEL, COLEGIO, PERIODO
    """
    es_nuevo = not os.path.exists(HISTORIAL_PATH)

    encabezado = [
        "FECHA",
        "NOMBRE_ESTUDIANTE",
        "CEDULA_ESTUDIANTE",
        "TELEFONO",
        "ESTADO",
        "NOMBRE_ACUDIENTE",
        "CEDULA_ACUDIENTE",
        "PLANILLA",
        "CHEQUE",
        "NIVEL",
        "COLEGIO",
        "PERIODO",
    ]

    fila = [
        datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        datos["NOMBRE_ESTUDIANTE"],
        datos["CEDULA_ESTUDIANTE"],
        telefono,
        estado,
        datos["NOMBRE_ACUDIENTE"],
        datos["CEDULA_ACUDIENTE"],
        datos["PLANILLA"],
        datos["CHEQUE"],
        datos["NIVEL"],
        datos["COLEGIO"],
        periodo,
    ]

    with open(HISTORIAL_PATH, "a", encoding="utf-8", newline="") as f:
        if es_nuevo:
            f.write(";".join(encabezado) + "\n")
        f.write(";".join(fila) + "\n")

def leer_historial():
    if not os.path.exists(HISTORIAL_PATH):
        return [], []
    with open(HISTORIAL_PATH, "r", encoding="utf-8") as f:
        lineas = [l.strip() for l in f.readlines() if l.strip()]
    if not lineas:
        return [], []
    encabezado = lineas[0].split(";")
    filas = [l.split(";") for l in lineas[1:]]
    return encabezado, filas

def actualizar_estado_reemplazo(cedula: str, cheques_seleccionados):
    """
    Marca como REEMPLAZO RECIBIDO los registros de esa cédula
    cuyos cheques estén en cheques_seleccionados.
    """
    if not os.path.exists(HISTORIAL_PATH):
        return None, []

    encabezado, filas = leer_historial()
    if not encabezado:
        return None, []

    idx_map = {name: i for i, name in enumerate(encabezado)}
    idx_ced = idx_map["CEDULA_ESTUDIANTE"]
    idx_est = idx_map["ESTADO"]
    idx_tel = idx_map["TELEFONO"]
    idx_chq = idx_map["CHEQUE"]

    telefono = None
    cheques_actualizados = []

    for r in filas:
        if len(r) < len(encabezado):
            continue
        if r[idx_ced] == cedula and r[idx_chq] in cheques_seleccionados:
            r[idx_est] = "REEMPLAZO RECIBIDO"
            cheques_actualizados.append(r[idx_chq])
            if telefono is None:
                telefono = r[idx_tel]

    if not cheques_actualizados:
        return None, []

    # Reescribir archivo
    with open(HISTORIAL_PATH, "w", encoding="utf-8", newline="") as f:
        f.write(";".join(encabezado) + "\n")
        for r in filas:
            f.write(";".join(r) + "\n")

    return telefono, cheques_actualizados

# =========================
# DOCX (NO TOCAR PLANTILLA, SOLO VARIABLES)
# =========================

def formatear_variables(doc, datos):
    """
    Solo toca los placeholders dentro de las tablas,
    reemplazando texto en los runs para NO borrar imágenes ni formato.
    """
    campos = [
        "CHEQUE", "PLANILLA", "NIVEL", "COLEGIO",
        "NOMBRE_ESTUDIANTE", "CEDULA_ESTUDIANTE",
        "NOMBRE_ACUDIENTE", "CEDULA_ACUDIENTE",
    ]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        original = r.text
                        if not original:
                            continue
                        for k in campos:
                            ph = f"{{{{{k}}}}}"
                            if ph in original:
                                original = original.replace(ph, datos[k])
                        if r.text != original:
                            r.text = original
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r.bold = True

def rellenar_docx(fila, acudiente, telefono):
    """
    Crea un Word a partir de una fila de la tabla y el acudiente.
    NO modifica la plantilla, solo rellena variables.
    """
    if not os.path.exists(PLANTILLA_PATH):
        raise FileNotFoundError("No se encontró plantilla.docx")

    doc = Document(PLANTILLA_PATH)

    datos = {
        "FECHA_ACTUAL": datetime.now().strftime("%d/%m/%Y"),
        "NOMBRE_ESTUDIANTE": fila["nombre"],
        "CEDULA_ESTUDIANTE": fila["cedula"],
        "NIVEL": fila["grado"],
        "COLEGIO": fila["centro"],
        "PLANILLA": fila["planilla"],
        "CHEQUE": fila["cheque"],
        "NOMBRE_ACUDIENTE": acudiente["NOMBRE_ACUDIENTE"],
        "CEDULA_ACUDIENTE": acudiente["CEDULA_ACUDIENTE"],
    }

    # Marcar periodo
    datos.update(periodo_a_checks(fila["periodo"]))

    # Reemplazar en párrafos normales (no tablas)
    for p in doc.paragraphs:
        for r in p.runs:
            original = r.text
            if not original:
                continue
            for k, v in datos.items():
                ph = f"{{{{{k}}}}}"
                if ph in original:
                    original = original.replace(ph, v)
            r.text = original

    # Reemplazar + formatear SOLO variables en tablas
    formatear_variables(doc, datos)

    # Nombre de archivo automático
    filename = f"{fila['cedula']}_{fila['periodo']}_{fila['cheque']}.docx"
    filename = filename.replace("/", "-")
    ruta = os.path.join(SALIDAS_DIR, filename)
    doc.save(ruta)

    # Guardar en historial
    guardar_en_historial(datos, telefono, fila["periodo"])

    return ruta

# =========================
# RUTAS WEB
# =========================

@app.route("/", methods=["GET", "POST"])
def index():
    texto = ""
    telefono = ""
    filas = []
    mensaje = ""

    if request.method == "POST":
        texto = request.form.get("texto", "")
        telefono = request.form.get("telefono", "")
        accion = request.form.get("accion")

        filas = parse_tabla_cheques(texto)
        acudiente = extraer_acudiente(texto)

        if not filas:
            mensaje = "No se detectaron filas en el texto."

        # Generar Word(s)
        if accion == "generar" and filas:
            seleccion = request.form.getlist("fila_idx")
            if not seleccion:
                flash("Debes seleccionar al menos un cheque.", "error")
            else:
                rutas = []
                for idx in seleccion:
                    try:
                        i = int(idx)
                    except ValueError:
                        continue
                    if 0 <= i < len(filas):
                        rutas.append(rellenar_docx(filas[i], acudiente, telefono))

                if rutas:
                    archivos = [os.path.basename(r) for r in rutas]
                    # Página que muestra enlaces y dispara descargas
                    return render_template("descargas.html", archivos=archivos)

                flash("No se pudo generar ningún documento.", "error")

    return render_template("index.html", texto=texto, telefono=telefono, filas=filas, mensaje=mensaje)

@app.route("/descargar/<filename>")
def descargar(filename):
    return send_from_directory(SALIDAS_DIR, filename, as_attachment=True)

@app.route("/historial")
def historial():
    encabezado, filas = leer_historial()
    return render_template("historial.html", encabezado=encabezado, filas=filas)

@app.route("/buscar", methods=["GET","POST"])
def buscar():
    resultados = []
    cedula = ""

    encabezado, filas = leer_historial()
    if request.method == "POST":
        cedula = request.form.get("cedula", "").strip()
        if cedula and encabezado:
            idx_map = {name:i for i,name in enumerate(encabezado)}
            for r in filas:
                if len(r) >= len(encabezado) and r[idx_map["CEDULA_ESTUDIANTE"]] == cedula:
                    resultados.append(r)

    return render_template("buscar.html", encabezado=encabezado, resultados=resultados, cedula=cedula)

@app.route("/reemplazo", methods=["GET","POST"])
def reemplazo():
    encabezado, filas = leer_historial()
    pendientes = []
    cedula = ""
    mensaje = ""

    if request.method == "POST" and encabezado:
        cedula = request.form.get("cedula","").strip()
        accion = request.form.get("accion")
        idx_map = {name:i for i,name in enumerate(encabezado)}

        if accion == "buscar":
            for r in filas:
                if (len(r) >= len(encabezado)
                        and r[idx_map["CEDULA_ESTUDIANTE"]] == cedula
                        and r[idx_map["ESTADO"]] != "REEMPLAZO RECIBIDO"):
                    pendientes.append(r)

        elif accion == "marcar":
            cheques_sel = request.form.getlist("cheque_sel")
            telefono, cheques_actualizados = actualizar_estado_reemplazo(cedula, cheques_sel)

            if not cheques_actualizados:
                mensaje = "No se pudo actualizar."
            else:
                mensaje = f"Teléfono: {telefono} | Cheques actualizados: {', '.join(cheques_actualizados)}"
                flash(mensaje, "ok")

            # recargar pendientes
            encabezado, filas = leer_historial()
            idx_map = {name:i for i,name in enumerate(encabezado)}
            for r in filas:
                if (len(r) >= len(encabezado)
                        and r[idx_map["CEDULA_ESTUDIANTE"]] == cedula
                        and r[idx_map["ESTADO"]] != "REEMPLAZO RECIBIDO"):
                    pendientes.append(r)

    return render_template("reemplazo.html",
                           encabezado=encabezado,
                           pendientes=pendientes,
                           cedula=cedula,
                           mensaje=mensaje)

if __name__ == "__main__":
    # Servidor web por red local
    app.run(host="0.0.0.0", port=5000, debug=True)
